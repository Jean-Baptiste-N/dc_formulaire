import re
import calendar
from datetime import datetime
from docx import Document

# ===========================================================================
# MARK: TEXT CLEANING FUNCTION
# ===========================================================================

def clean_text(text):
    """
    Nettoie le texte en supprimant les caractères indésirables.
    Utile pour nettoyer les données collées depuis Word ou d'autres sources.
    """
    if not isinstance(text, str):
        return text

    # Enlever les guillemets littéraux au début et fin (notamment du JSON avec double-encoding)
    text = re.sub(r'^["\']|["\']$', '', text)

    # Supprimer caractères non-alphanumériques au début et fin
    text = re.sub(r'^\W+', '', text)
    text = re.sub(r'\W+$', '', text)

    return text.strip()

# ============================================================================
# MARK: DOCX CLEANING FUNCTION
# ============================================================================

def remove_empty_paragraphs(docx_path):
    """Supprime les bullets vides ET les paragraphes vides dans les tables (colonne 1 seulement)"""
    doc = Document(docx_path)

    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Première passe: nettoyer les runs vides
    for paragraph in doc.paragraphs:
        runs_to_remove = []
        has_page_break = False

        # Vérifier si le paragraphe contient un page break
        for run in paragraph.runs:
            br_elem = run._element.find('{' + w_ns + '}br')
            if br_elem is not None and br_elem.get('{' + w_ns + '}type') == 'page':
                has_page_break = True
                break

        # Supprimer les runs vides SAUF s'il y a un page break ou une image
        if not has_page_break:
            for run in paragraph.runs:
                if not run.text.strip():
                    # Ne pas supprimer si le run contient une image
                    has_drawing = run._element.find('{' + w_ns + '}drawing') is not None
                    if not has_drawing:
                        runs_to_remove.append(run._element)

        for run_elem in reversed(runs_to_remove):
            run_elem.getparent().remove(run_elem)

    # Deuxième passe: supprimer les bullets vides
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if not text:
            style_name = paragraph.style.name  # type: ignore[union-attr]
            # Supprimer les bullets vides (DC_*_bullet)
            if '_bullet' in style_name:  # type: ignore[operator]
                paragraphs_to_remove.append(paragraph._element)

    for p_element in reversed(paragraphs_to_remove):
        p_element.getparent().remove(p_element)

    doc.save(docx_path)

# ============================================================================
# MARK: DATE PARSING & SORTING FUNCTIONS
# ============================================================================

def parse_date_to_end_date(date_str):
    """
    Parse une date au format texte et retourne la date de fin.
    Gère plusieurs formats:
    - "09/2025 – 03/2026" -> 2026-03-31
    - "2014-2016" -> 2016-12-31
    - "2019" -> 2019-12-31
    - "11-2024 - 05/2026" -> 2026-05-31
    - "09/2018 – 04/2019" -> 2019-04-30
    - "09/2025" -> 2025-09-30
    - "11-2024" -> 2024-11-30
    - "depuis 11/2021" -> 2021-11-30 (texte + date, prend la dernière date)
    - "jusqu'à 09/2025" -> 2025-09-30
    - "de 2019 à 2021" -> 2021-12-31

    Retourne un tuple (date_obj, date_str) pour le tri.
    """
    if not date_str or not isinstance(date_str, str):
        return (datetime.min, "")

    date_str = date_str.strip()

    # Normaliser les tirets: – et — et - vers un format standard
    normalized = re.sub(r'[–—-]+', '-', date_str)  # Remplacer tous les tirets par un seul tiret
    normalized = re.sub(r'\s+', ' ', normalized)    # Normaliser les espaces

    # Helper pour obtenir le dernier jour du mois
    def get_last_day_of_month(year, month):
        """Retourne le dernier jour du mois donné."""
        last_day = calendar.monthrange(year, month)[1]
        return datetime(year, month, last_day)

    # Format: "MM/YYYY – MM/YYYY" ou "MM/YYYY - MM/YYYY"
    month_year_range = re.search(
        r'(\d{1,2})/(\d{4})\s*(?:-|–|—)\s*(\d{1,2})/(\d{4})',
        normalized
    )
    if month_year_range:
        _, _, end_month, end_year = month_year_range.groups()
        try:
            end_month = int(end_month)
            end_year = int(end_year)
            date_obj = get_last_day_of_month(end_year, end_month)
            return (date_obj, date_str)
        except (ValueError, OverflowError):
            pass

    # Format: "YYYY-YYYY" (ex: 2014-2016)
    year_range = re.search(r'(\d{4})\s*(?:-|–|—)\s*(\d{4})', normalized)
    if year_range:
        start_year, end_year = year_range.groups()
        try:
            end_year = int(end_year)
            date_obj = datetime(end_year, 12, 31)
            return (date_obj, date_str)
        except (ValueError, OverflowError):
            pass

    # Format: "MM-YYYY - MM/YYYY" (ex: 11-2024 - 05/2026) - Plage de deux dates
    month_dash_year_range = re.search(
        r'(\d{1,2})(?:-|/)?(\d{4})\s*(?:-|–|—)\s*(\d{1,2})(?:-|/)?(\d{4})',
        normalized
    )
    if month_dash_year_range:
        _, _, end_month, end_year = month_dash_year_range.groups()
        try:
            end_month = int(end_month)
            end_year = int(end_year)
            date_obj = get_last_day_of_month(end_year, end_month)
            return (date_obj, date_str)
        except (ValueError, OverflowError):
            pass

    # Format: "MM/YYYY" seul (ex: 09/2025)
    month_year_single = re.search(r'^(\d{1,2})/(\d{4})$', normalized.strip())
    if month_year_single:
        try:
            month = int(month_year_single.group(1))
            year = int(month_year_single.group(2))
            date_obj = get_last_day_of_month(year, month)
            return (date_obj, date_str)
        except (ValueError, OverflowError):
            pass

    # Format: "MM-YYYY" seul (ex: 11-2024)
    month_dash_year_single = re.search(r'^(\d{1,2})-(\d{4})$', normalized.strip())
    if month_dash_year_single:
        try:
            month = int(month_dash_year_single.group(1))
            year = int(month_dash_year_single.group(2))
            date_obj = get_last_day_of_month(year, month)
            return (date_obj, date_str)
        except (ValueError, OverflowError):
            pass

    # Format: juste une année "YYYY" (ex: 2019)
    year_only = re.search(r'^(\d{4})$', normalized.strip())
    if year_only:
        try:
            year = int(year_only.group(1))
            date_obj = datetime(year, 12, 31)
            return (date_obj, date_str)
        except (ValueError, OverflowError):
            pass

    # Format: texte + date (ex: "depuis 11/2021", "depuis 09/2025", "depuis 2019")
    # Extraire la DERNIÈRE date trouvée dans la chaîne
    all_dates_with_month = []
    all_dates_year_only = []

    # Chercher tous les "MM/YYYY"
    for match in re.finditer(r'(\d{1,2})/(\d{4})', normalized):
        month = int(match.group(1))
        year = int(match.group(2))
        if 1 <= month <= 12:
            all_dates_with_month.append((match.start(), get_last_day_of_month(year, month)))

    # Chercher tous les "MM-YYYY"
    for match in re.finditer(r'(\d{1,2})-(\d{4})', normalized):
        month = int(match.group(1))
        year = int(match.group(2))
        if 1 <= month <= 12:
            all_dates_with_month.append((match.start(), get_last_day_of_month(year, month)))

    # Si on a trouvé des dates avec mois, prendre la DERNIÈRE
    if all_dates_with_month:
        try:
            # Trier par position pour prendre la dernière (la plus à droite)
            all_dates_with_month.sort(key=lambda x: x[0], reverse=True)
            date_obj = all_dates_with_month[0][1]
            return (date_obj, date_str)
        except (ValueError, OverflowError):
            pass

    # Sinon, chercher les "YYYY" seules (4 chiffres seuls, pas dans une plage)
    for match in re.finditer(r'\b(\d{4})\b', normalized):
        year = int(match.group(1))
        all_dates_year_only.append((match.start(), datetime(year, 12, 31)))

    # Si on a trouvé des années seules, prendre la DERNIÈRE
    if all_dates_year_only:
        try:
            # Trier par position pour prendre la dernière (la plus à droite)
            all_dates_year_only.sort(key=lambda x: x[0], reverse=True)
            date_obj = all_dates_year_only[0][1]
            return (date_obj, date_str)
        except (ValueError, OverflowError):
            pass

    # Si aucun format ne correspond, retourner une date minimale
    return (datetime.min, date_str)


def sort_items_by_date_desc(items, date_field="date"):
    """
    Trie une liste d'items (formations, certifications, xp_pro) en anti-chronologique.
    Les items avec les dates les plus récentes apparaissent en premier.

    Args:
        items: Liste des items à trier
        date_field: Nom du champ contenant la date (par défaut "date")

    Returns:
        Liste triée en anti-chronologique
    """
    if not isinstance(items, list):
        return items

    # Créer une copie pour ne pas modifier la liste originale
    sorted_items = items.copy()

    # Trier en mettant les dates les plus récentes en premier (descending)
    sorted_items.sort(
        key=lambda item: parse_date_to_end_date(item.get(date_field, ""))[0],
        reverse=True  # Anti-chronologique: plus récent en premier
    )

    return sorted_items


def sort_dossier_items(dossier):
    """
    Trie tous les items du dossier par date (formations, certifications).
    NB: xp_pro ne sont PAS triés car ils ont un ordre personnel défini par l'utilisateur via order_index.
    Modifie le dossier in-place et le retourne.

    Args:
        dossier: Le dossier complet en dictionnaire JSON

    Returns:
        Le dossier modifié avec les listes triées
    """
    if not isinstance(dossier, dict):
        return dossier

    # Trier les formations
    if "formations" in dossier and isinstance(dossier["formations"], list):
        dossier["formations"] = sort_items_by_date_desc(
            dossier["formations"],
            date_field="date"
        )

    # Trier les certifications
    if "certifications" in dossier and isinstance(dossier["certifications"], list):
        dossier["certifications"] = sort_items_by_date_desc(
            dossier["certifications"],
            date_field="date"
        )

    # ⚠️ xp_pro ne sont PAS triés - l'utilisateur les ordonne manuellement via order_index
    # Les xp_pro sont triés par order_index dans les vues si nécessaire

    return dossier


def sort_xp_pro_by_order_index(xp_pro_list):
    """
    Trie les xp_pro par leur order_index en ordre croissant (index 1 en haut).
    Les xp_pro sans order_index sont placés à la fin.

    Args:
        xp_pro_list: Liste des xp_pro à trier

    Returns:
        Liste triée par order_index (croissant)
    """
    if not isinstance(xp_pro_list, list):
        return xp_pro_list

    # Créer une copie
    sorted_list = xp_pro_list.copy()

    # Trier: d'abord par order_index croissant (les items sans sont à la fin)
    def get_order_index(item):
        """Retourne l'order_index ou un très grand nombre s'il n'existe pas."""
        order_idx = item.get("order_index")
        if order_idx is None:
            return float('inf')
        try:
            return int(order_idx)
        except (ValueError, TypeError):
            return float('inf')

    sorted_list.sort(key=get_order_index)
    return sorted_list


# ============================================================================
# MARK: NORMALIZATION FUNCTIONS - Add UUID and order_index
# ============================================================================

def normalize_item_with_id_and_order(item, item_type="formation"):
    """
    Normalise un item en ajoutant un UUID et order_index s'ils manquent.

    Args:
        item: L'item à normaliser (formation, certification, ou langue)
        item_type: Type d'item ('formation', 'certification', 'langue')

    Returns:
        L'item normalisé avec UUID et order_index
    """
    import uuid as uuid_module

    if not isinstance(item, dict):
        return item

    # Ajouter UUID si absent
    if "id" not in item or not item["id"]:
        item["id"] = str(uuid_module.uuid4())

    # Ajouter order_index si absent (sera recalculé après)
    if "order_index" not in item:
        item["order_index"] = 1

    return item


def normalize_items_list_with_order(items, item_type="formation"):
    """
    Normalise une liste d'items en ajoutant UUID et recalculant order_index par date.

    Args:
        items: Liste des items à normaliser
        item_type: Type d'item ('formation', 'certification', 'langue')

    Returns:
        Liste normalisée avec UUID et order_index recalculés
    """
    import uuid as uuid_module

    if not isinstance(items, list):
        return items

    # D'abord, ajouter les UUIDs manquants
    for item in items:
        if "id" not in item or not item.get("id"):
            item["id"] = str(uuid_module.uuid4())

    # Ensuite, trier par date (anti-chronologique) et assigner order_index
    if item_type in ["formation", "certification"]:
        # Trier par date décroissante (dates récentes en premier)
        items_with_dates = []
        for i, item in enumerate(items):
            date_obj, _ = parse_date_to_end_date(item.get("date", ""))
            items_with_dates.append((date_obj, i, item))

        # Trier par date décroissante
        items_with_dates.sort(key=lambda x: x[0], reverse=True)

        # Réassigner order_index et réordonner la liste
        normalized_items = []
        for order_idx, (_, _, item) in enumerate(items_with_dates, start=1):
            item["order_index"] = order_idx
            normalized_items.append(item)

        return normalized_items
    elif item_type == "langue":
        # Pour les langues, garder l'ordre et assigner order_index séquentiellement
        for order_idx, item in enumerate(items, start=1):
            item["order_index"] = order_idx
        return items

    return items
