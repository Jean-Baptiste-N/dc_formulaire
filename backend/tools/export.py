import argparse
import json

import docxtpl
from docx import Document


def remove_empty_paragraphs(docx_path):
    """Supprime les bullets vides ET les paragraphes vides dans les tables (colonne 1 seulement)"""
    doc = Document(docx_path)

    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Première passe: nettoyer les runs vides
    for paragraph in doc.paragraphs:
        runs_to_remove = []
        has_page_break = False

        # Vérifier si le paragraphe contient un page break
        for run in paragraph.runs:
            br_elem = run._element.find('{' + w_ns + '}br')
            if br_elem is not None and br_elem.get('{' + w_ns + '}type') == 'page':
                has_page_break = True
                break

        # Supprimer les runs vides SAUF s'il y a un page break ou une image
        if not has_page_break:
            for run in paragraph.runs:
                if not run.text.strip():
                    # Ne pas supprimer si le run contient une image
                    has_drawing = run._element.find('{' + w_ns + '}drawing') is not None
                    if not has_drawing:
                        runs_to_remove.append(run._element)

        for run_elem in reversed(runs_to_remove):
            run_elem.getparent().remove(run_elem)

    # Deuxième passe: supprimer les bullets vides
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if not text:
            style_name = paragraph.style.name  # type: ignore[union-attr]
            # Supprimer les bullets vides (DC_*_bullet)
            if '_bullet' in style_name:  # type: ignore[operator]
                paragraphs_to_remove.append(paragraph._element)

    for p_element in reversed(paragraphs_to_remove):
        p_element.getparent().remove(p_element)

    doc.save(docx_path)

def main(input_file="inputs/DC_JNZ_2026_uuid.json", output_path="outputs/exemple.docx"):
    with open(input_file, encoding='utf-8') as json_file:
        json_data = json.load(json_file)

    template = docxtpl.DocxTemplate('templates_docx/template_jinja.docx')
    template.render(json_data)
    template.save(output_path)
    print(f"✅ Document généré : {output_path}")

    # Nettoyage post-génération
    remove_empty_paragraphs(output_path)
    print(f"✅ Nettoyage terminé : {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Export formulaire en DOCX')
    parser.add_argument('-i', '--input', default='inputs/DC_JNZ_2026_uuid.json', help='Fichier JSON source (défaut: inputs/DC_JNZ_2026_uuid.json)')
    parser.add_argument('-o', '--output', default='outputs/exemple.docx', help='Nom du fichier de sortie (défaut: outputs/exemple.docx)')
    args = parser.parse_args()

    main(args.input, args.output)

